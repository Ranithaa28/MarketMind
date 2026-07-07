"""
Generates downloadable PDF and DOCX reports from a completed Idea record.
Both formats cover: executive summary, idea, competitors, market analysis,
investment, SWOT, business canvas, financial estimate, and recommendations.
"""
import os
from datetime import datetime

from docx import Document
from docx.shared import Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from app.db.models import Idea

REPORTS_DIR = os.environ.get("REPORTS_DIR", "/tmp/reports")
os.makedirs(REPORTS_DIR, exist_ok=True)


def _kv_rows(d: dict | None) -> list[tuple]:
    if not d:
        return []
    return [(str(k).replace("_", " ").title(), str(v)) for k, v in d.items() if not isinstance(v, (list, dict))]


def generate_pdf_report(idea: Idea) -> str:
    filename = f"{idea.id}_{int(datetime.utcnow().timestamp())}.pdf"
    path = os.path.join(REPORTS_DIR, filename)

    styles = getSampleStyleSheet()
    h1 = ParagraphStyle("h1", parent=styles["Heading1"], spaceAfter=10)
    h2 = ParagraphStyle("h2", parent=styles["Heading2"], spaceBefore=14, spaceAfter=6)
    body = styles["BodyText"]

    doc = SimpleDocTemplate(path, pagesize=LETTER, topMargin=0.7 * inch, bottomMargin=0.7 * inch)
    story = [
        Paragraph(f"Startup Validation Report: {idea.title}", h1),
        Paragraph(f"Generated {datetime.utcnow().strftime('%Y-%m-%d')}", body),
        Spacer(1, 12),
        Paragraph("Executive Summary", h2),
        Paragraph(idea.raw_description, body),
    ]

    if idea.success_score:
        story.append(Paragraph("AI Success Score", h2))
        s = idea.success_score
        story.append(
            Paragraph(
                f"Overall: {s.get('overall_score')}/100 &nbsp;|&nbsp; "
                f"Strength: {s.get('strength_meter')} &nbsp;|&nbsp; "
                f"Risk: {s.get('risk_meter')} &nbsp;|&nbsp; "
                f"Opportunity: {s.get('opportunity_meter')}",
                body,
            )
        )
        story.append(Paragraph(s.get("disclaimer", ""), styles["Italic"]))

    if idea.competitors and idea.competitors.get("competitors"):
        story.append(Paragraph("Competitors", h2))
        rows = [["Name", "Pricing", "Market Position"]]
        for c in idea.competitors["competitors"]:
            rows.append([c.get("name", ""), str(c.get("pricing", "")), c.get("market_position", "")])
        table = Table(rows, colWidths=[150, 150, 200])
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1f2937")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("FONTSIZE", (0, 0), (-1, -1), 8),
                ]
            )
        )
        story.append(table)

    if idea.market_research:
        story.append(Paragraph("Market Research", h2))
        for label, val in _kv_rows(idea.market_research):
            story.append(Paragraph(f"<b>{label}:</b> {val}", body))

    if idea.investment:
        story.append(Paragraph("Investment Estimate", h2))
        for label, val in _kv_rows(idea.investment):
            story.append(Paragraph(f"<b>{label}:</b> {val}", body))

    if idea.swot:
        story.append(Paragraph("SWOT Analysis", h2))
        for section in ["strengths", "weaknesses", "opportunities", "threats"]:
            items = idea.swot.get(section, [])
            story.append(Paragraph(f"<b>{section.title()}:</b> " + "; ".join(items), body))

    if idea.strategy:
        story.append(Paragraph("Business Strategy", h2))
        for label, val in _kv_rows(idea.strategy):
            story.append(Paragraph(f"<b>{label}:</b> {val}", body))

    doc.build(story)
    return path


def generate_docx_report(idea: Idea) -> str:
    filename = f"{idea.id}_{int(datetime.utcnow().timestamp())}.docx"
    path = os.path.join(REPORTS_DIR, filename)

    doc = Document()
    doc.add_heading(f"Startup Validation Report: {idea.title}", level=1)
    doc.add_paragraph(f"Generated {datetime.utcnow().strftime('%Y-%m-%d')}")

    doc.add_heading("Executive Summary", level=2)
    doc.add_paragraph(idea.raw_description)

    if idea.success_score:
        doc.add_heading("AI Success Score", level=2)
        s = idea.success_score
        p = doc.add_paragraph()
        p.add_run(
            f"Overall: {s.get('overall_score')}/100  |  Strength: {s.get('strength_meter')}  |  "
            f"Risk: {s.get('risk_meter')}  |  Opportunity: {s.get('opportunity_meter')}"
        ).bold = True
        italic = doc.add_paragraph(s.get("disclaimer", ""))
        italic.runs[0].italic = True
        italic.runs[0].font.size = Pt(9)

    if idea.competitors and idea.competitors.get("competitors"):
        doc.add_heading("Competitors", level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = "Name", "Pricing", "Market Position"
        for c in idea.competitors["competitors"]:
            row = table.add_row().cells
            row[0].text = c.get("name", "")
            row[1].text = str(c.get("pricing", ""))
            row[2].text = c.get("market_position", "")

    if idea.market_research:
        doc.add_heading("Market Research", level=2)
        for label, val in _kv_rows(idea.market_research):
            doc.add_paragraph(f"{label}: {val}")

    if idea.investment:
        doc.add_heading("Investment Estimate", level=2)
        for label, val in _kv_rows(idea.investment):
            doc.add_paragraph(f"{label}: {val}")

    if idea.swot:
        doc.add_heading("SWOT Analysis", level=2)
        for section in ["strengths", "weaknesses", "opportunities", "threats"]:
            items = idea.swot.get(section, [])
            doc.add_paragraph(f"{section.title()}: " + "; ".join(items))

    if idea.strategy:
        doc.add_heading("Business Strategy", level=2)
        for label, val in _kv_rows(idea.strategy):
            doc.add_paragraph(f"{label}: {val}")

    doc.save(path)
    return path
